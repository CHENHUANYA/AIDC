from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
INPUT = ROOT / "docs" / "期末計畫書撰寫摘要.md"
OUTPUT = ROOT / "docs" / "期末計畫書撰寫摘要.docx"

INLINE_CODE = re.compile(r"`([^`]+)`")
STRONG = re.compile(r"\*\*([^*]+)\*\*")
ASCII_RUN = re.compile(r"([A-Za-z0-9][A-Za-z0-9_\-./:()&+, ]*[A-Za-z0-9)]|[A-Za-z0-9])")


def clean_text(text: str) -> str:
    text = INLINE_CODE.sub(lambda match: match.group(1), text)
    text = STRONG.sub(r"\1", text)
    return text.strip()


def set_run_font(run, *, size: int = 12, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "DFKai-SB")
    rfonts.set(qn("w:cs"), "Times New Roman")


def add_mixed_text(paragraph, text: str, *, size: int = 12, bold: bool = False) -> None:
    text = clean_text(text)
    pos = 0
    for match in ASCII_RUN.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=size, bold=bold)
        run = paragraph.add_run(match.group(0))
        set_run_font(run, size=size, bold=bold)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size, bold=bold)


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "DFKai-SB")
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)


def format_paragraph(paragraph, *, align=None, first_line: bool = False) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    if first_line:
        fmt.first_line_indent = Pt(24)
    if align is not None:
        paragraph.alignment = align


def is_table_start(lines: list[str], idx: int) -> bool:
    if idx + 1 >= len(lines):
        return False
    return lines[idx].strip().startswith("|") and re.match(r"^\s*\|?\s*:?-{3,}", lines[idx + 1])


def split_table_row(line: str) -> list[str]:
    return [clean_text(cell) for cell in line.strip().strip("|").split("|")]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            text = row[c_idx] if c_idx < len(row) else ""
            para = cell.paragraphs[0]
            format_paragraph(para)
            add_mixed_text(para, text, size=10 if col_count > 4 else 11, bold=(r_idx == 0))
            if r_idx == 0:
                set_cell_shading(cell, "F2F2F2")
    spacer = doc.add_paragraph()
    format_paragraph(spacer)


def add_footer_page_number(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, size=10)


def main() -> None:
    doc = Document()
    setup_document(doc)
    add_footer_page_number(doc)

    lines = INPUT.read_text(encoding="utf-8").splitlines()
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        if not paragraph_buffer:
            return
        text = " ".join(item.strip() for item in paragraph_buffer if item.strip())
        paragraph_buffer.clear()
        if text:
            para = doc.add_paragraph()
            format_paragraph(para, first_line=True)
            add_mixed_text(para, text, size=12)

    idx = 0
    while idx < len(lines):
        raw = lines[idx]
        stripped = raw.strip()

        if not stripped:
            flush_paragraph()
            idx += 1
            continue

        if is_table_start(lines, idx):
            flush_paragraph()
            rows = [split_table_row(lines[idx])]
            idx += 2
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                rows.append(split_table_row(lines[idx]))
                idx += 1
            add_table(doc, rows)
            continue

        if stripped.startswith("# "):
            flush_paragraph()
            para = doc.add_paragraph()
            format_paragraph(para, align=WD_ALIGN_PARAGRAPH.CENTER)
            add_mixed_text(para, stripped[2:], size=16, bold=True)
        elif stripped.startswith("## "):
            flush_paragraph()
            heading = stripped[3:].strip()
            if heading.startswith("(七)"):
                para = doc.add_paragraph()
                para.add_run().add_break(WD_BREAK.PAGE)
            para = doc.add_paragraph()
            format_paragraph(para)
            add_mixed_text(para, heading, size=14, bold=True)
        elif stripped.startswith("### "):
            flush_paragraph()
            para = doc.add_paragraph()
            format_paragraph(para)
            add_mixed_text(para, stripped[4:], size=12, bold=True)
        elif re.match(r"^\d+\.\s+", stripped):
            flush_paragraph()
            para = doc.add_paragraph()
            format_paragraph(para)
            add_mixed_text(para, stripped, size=12)
        elif stripped.startswith("- "):
            flush_paragraph()
            para = doc.add_paragraph()
            format_paragraph(para)
            add_mixed_text(para, "• " + stripped[2:], size=12)
        else:
            paragraph_buffer.append(stripped)
        idx += 1

    flush_paragraph()
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
